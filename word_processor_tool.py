# pip install python-docx pywin32
import tkinter as tk
from pathlib import Path
from tkinter import filedialog, messagebox, ttk
import os
import docx
import win32com.client as win32
import threading


def convert_doc_to_docx(doc_file_path):
    """
    .docファイルを.docxファイルに変換する

    Args:
        doc_file_path (str): 変換する.docファイルのパス

    Returns:
        str: 変換された.docxファイルのパス
    """
    try:
        # パスを文字列に変換して絶対パスを取得
        doc_file_path = os.path.abspath(str(doc_file_path))

        word = win32.Dispatch('Word.Application')
        word.Visible = False  # Wordを非表示にする

        # ファイルを開く
        doc = word.Documents.Open(doc_file_path)

        # .docxファイルのパスを生成
        docx_file_path = os.path.splitext(doc_file_path)[0] + '.docx'

        # .docファイルを.docx形式で保存します (FileFormat=16: .docx format)
        doc.SaveAs2(docx_file_path, FileFormat=16)  # 16: .docx format
        doc.Close()

        word.Quit()
        return docx_file_path
    except Exception as e:
        print(f"ファイル '{doc_file_path}' の変換中にエラーが発生しました: {e}")
        return None


def batch_convert_docs_in_folder(folder_path, delete_original=False, progress_callback=None, status_callback=None):
    """
    フォルダ内の全ての.docファイルを.docxファイルに変換する

    Args:
        folder_path (str): .docファイルを含むフォルダのパス
        delete_original (bool): 変換後に元の.docファイルを削除するかどうか
        progress_callback (function): 進捗状況を報告するコールバック関数
        status_callback (function): 状態メッセージを報告するコールバック関数

    Returns:
        tuple: (処理されたファイル数, 変換されたファイル数)
    """
    # フォルダ内のすべての.docファイルを取得
    doc_files = list(Path(folder_path).glob("**/*.doc"))
    total_files = len(doc_files)
    converted_files = 0

    if status_callback:
        status_callback(f"合計 {total_files} 件の.docファイルが見つかりました。")

    if total_files == 0:
        if status_callback:
            status_callback("変換対象の.docファイルが見つかりませんでした。")
        return 0, 0

    # 各.docファイルを処理
    for i, doc_file_path in enumerate(doc_files):
        try:
            if status_callback:
                status_callback(f"変換中: {doc_file_path}")

            docx_path = convert_doc_to_docx(str(doc_file_path))
            if docx_path:
                converted_files += 1
                if status_callback:
                    status_callback(f"変換完了: {docx_path}")

                # オリジナルのdocファイルを削除
                if delete_original:
                    try:
                        os.remove(str(doc_file_path))
                        if status_callback:
                            status_callback(f"元のファイルを削除しました: {doc_file_path}")
                    except Exception as e:
                        if status_callback:
                            status_callback(f"元のファイル削除中にエラー: {e}")

            if progress_callback:
                progress_callback((i + 1) / total_files * 100)

        except Exception as e:
            if status_callback:
                status_callback(f"エラー: {e}")

    return total_files, converted_files


def replace_text_in_word_files(directory_path, old_text, new_text, preserve_formatting=True, progress_callback=None,
                               status_callback=None):
    """
    指定されたディレクトリ内のすべてのWordファイルで特定のテキストを置換します。

    Args:
        directory_path (str): Wordファイルを含むディレクトリのパス
        old_text (str): 置換対象のテキスト
        new_text (str): 置換後のテキスト
        preserve_formatting (bool): 書式を保持するかどうか
        progress_callback (function): 進捗状況を報告するコールバック関数
        status_callback (function): 状態メッセージを報告するコールバック関数

    Returns:
        tuple: (処理されたファイル数, 変更されたファイル数)
    """
    # ディレクトリ内のすべてのWordファイルを取得
    word_files = list(Path(directory_path).glob("**/*.docx"))
    total_files = len(word_files)
    modified_files = 0

    if status_callback:
        status_callback(f"合計 {total_files} 件の.docxファイルが見つかりました。")

    if total_files == 0:
        if status_callback:
            status_callback("処理対象の.docxファイルが見つかりませんでした。")
        if progress_callback:
            progress_callback(100)  # 処理するファイルがない場合、プログレスバーを100%にする
        return 0, 0

    # 各Wordファイルを処理
    for i, file_path in enumerate(word_files):
        try:
            if status_callback:
                status_callback(f"処理中 ({i + 1}/{total_files}): {file_path}")

            # Wordファイルを開く
            doc = docx.Document(str(file_path))

            # 置換が行われたかを追跡
            replacement_occurred = False

            if preserve_formatting:
                # 書式を保持して置換する（高度な方法）
                # すべての段落内のテキストを置換
                for paragraph in doc.paragraphs:
                    if old_text in paragraph.text:
                        replacement_occurred = True

                        # 段落内の各Runを処理（書式を保持）
                        runs = paragraph.runs
                        for i, run in enumerate(runs):
                            if old_text in run.text:
                                # 現在のRunの書式を保持しながら置換
                                run.text = run.text.replace(old_text, new_text)

                # すべての表内のテキストを置換
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if old_text in paragraph.text:
                                    replacement_occurred = True

                                    # 段落内の各Runを処理（書式を保持）
                                    runs = paragraph.runs
                                    for i, run in enumerate(runs):
                                        if old_text in run.text:
                                            # 現在のRunの書式を保持しながら置換
                                            run.text = run.text.replace(old_text, new_text)
            else:
                # 簡易置換（書式は保持されない）
                # すべての段落内のテキストを置換
                for paragraph in doc.paragraphs:
                    if old_text in paragraph.text:
                        replacement_occurred = True
                        text = paragraph.text
                        paragraph.text = text.replace(old_text, new_text)

                # すべての表内のテキストを置換
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if old_text in paragraph.text:
                                    replacement_occurred = True
                                    text = paragraph.text
                                    paragraph.text = text.replace(old_text, new_text)

            # 変更があった場合のみファイルを保存
            if replacement_occurred:
                doc.save(str(file_path))
                modified_files += 1
                if status_callback:
                    status_callback(f"置換完了 ({i + 1}/{total_files}): {file_path}")

            if progress_callback:
                progress_callback((i + 1) / total_files * 100)

        except Exception as e:
            if status_callback:
                status_callback(f"エラー: ファイル '{file_path}' の処理中にエラーが発生しました: {e}")

    # 最終的に確実にプログレスバーを100%にする
    if progress_callback:
        progress_callback(100)

    return total_files, modified_files


def select_directory():
    """
    ディレクトリ選択ダイアログを表示し、選択されたディレクトリパスを返します。
    """
    directory_path = filedialog.askdirectory(title="Wordファイルを含むディレクトリを選択")
    return directory_path


def create_gui():
    """
    GUIアプリケーションを作成して表示します。
    """
    # ルートウィンドウを作成
    root = tk.Tk()
    root.title("Wordファイル処理ツール")
    root.geometry("550x400")
    root.resizable(True, True)

    # タブコントロールを作成
    tab_control = ttk.Notebook(root)

    # タブを作成
    tab_convert = ttk.Frame(tab_control)
    tab_replace = ttk.Frame(tab_control)

    # タブをタブコントロールに追加
    tab_control.add(tab_convert, text='Doc→Docx変換')
    tab_control.add(tab_replace, text='テキスト置換')

    # タブコントロールをウィンドウに配置
    tab_control.pack(expand=1, fill="both")

    # 変数を定義
    convert_dir_var = tk.StringVar()
    replace_dir_var = tk.StringVar()
    old_text_var = tk.StringVar()
    new_text_var = tk.StringVar()
    status_var = tk.StringVar()
    status_var.set("待機中...")
    delete_original_var = tk.BooleanVar(value=False)
    preserve_formatting_var = tk.BooleanVar(value=True)

    # ----------------------
    # Doc→Docx変換タブの内容
    # ----------------------
    convert_frame = ttk.Frame(tab_convert, padding=20)
    convert_frame.pack(fill=tk.BOTH, expand=True)

    # ディレクトリ選択部分
    convert_dir_frame = ttk.Frame(convert_frame)
    convert_dir_frame.pack(fill=tk.X, pady=(0, 10))

    ttk.Label(convert_dir_frame, text="ディレクトリ:").pack(side=tk.LEFT)
    ttk.Entry(convert_dir_frame, textvariable=convert_dir_var, width=40).pack(side=tk.LEFT, padx=(5, 5), fill=tk.X,
                                                                              expand=True)

    def browse_convert_directory():
        directory_path = select_directory()
        if directory_path:
            convert_dir_var.set(directory_path)

    ttk.Button(convert_dir_frame, text="参照...", command=browse_convert_directory).pack(side=tk.LEFT)

    # 元ファイル削除オプション
    delete_option_frame = ttk.Frame(convert_frame)
    delete_option_frame.pack(fill=tk.X, pady=(0, 10))

    ttk.Checkbutton(delete_option_frame, text="変換後に元の.docファイルを削除する", variable=delete_original_var).pack(
        anchor=tk.W)

    # 実行ボタン
    convert_button_frame = ttk.Frame(convert_frame)
    convert_button_frame.pack(pady=(10, 20))

    # 進捗バー
    progress_frame = ttk.Frame(convert_frame)
    progress_frame.pack(fill=tk.X, pady=(0, 10))
    progress_bar = ttk.Progressbar(progress_frame, orient="horizontal", length=300, mode="determinate")
    progress_bar.pack(fill=tk.X)

    # ステータス表示
    status_frame = ttk.Frame(convert_frame)
    status_frame.pack(fill=tk.X)
    status_label = ttk.Label(status_frame, textvariable=status_var, wraplength=500)
    status_label.pack(fill=tk.X)

    def update_progress(value):
        progress_bar["value"] = value
        root.update_idletasks()

    def update_status(message):
        status_var.set(message)
        root.update_idletasks()

    def on_convert():
        directory_path = convert_dir_var.get()

        if not directory_path:
            messagebox.showerror("エラー", "ディレクトリを選択してください。")
            return

        # UIを無効化
        convert_button["state"] = "disabled"
        progress_bar["value"] = 0
        status_var.set("処理を開始します...")

        def process_thread():
            try:
                # 処理を実行
                processed_files, converted_files = batch_convert_docs_in_folder(
                    directory_path,
                    delete_original=delete_original_var.get(),
                    progress_callback=update_progress,
                    status_callback=update_status
                )

                # 結果を表示
                messagebox.showinfo("処理完了",
                                    f"処理が完了しました。\n"
                                    f"処理されたファイル数: {processed_files}\n"
                                    f"変換されたファイル数: {converted_files}")

                status_var.set(f"完了しました。処理ファイル数: {processed_files}, 変換ファイル数: {converted_files}")
            except Exception as e:
                messagebox.showerror("エラー", f"処理中にエラーが発生しました: {e}")
                status_var.set(f"エラー: {e}")
            finally:
                # UIを有効化
                root.after(0, lambda: convert_button.config(state="normal"))

        # 別スレッドで処理を実行
        threading.Thread(target=process_thread, daemon=True).start()

    convert_button = ttk.Button(convert_button_frame, text="変換実行", command=on_convert, width=15)
    convert_button.pack()

    # ----------------------
    # テキスト置換タブの内容
    # ----------------------
    replace_frame = ttk.Frame(tab_replace, padding=20)
    replace_frame.pack(fill=tk.BOTH, expand=True)

    # ディレクトリ選択部分
    replace_dir_frame = ttk.Frame(replace_frame)
    replace_dir_frame.pack(fill=tk.X, pady=(0, 10))

    ttk.Label(replace_dir_frame, text="ディレクトリ:").pack(side=tk.LEFT)
    ttk.Entry(replace_dir_frame, textvariable=replace_dir_var, width=40).pack(side=tk.LEFT, padx=(5, 5), fill=tk.X,
                                                                              expand=True)

    def browse_replace_directory():
        directory_path = select_directory()
        if directory_path:
            replace_dir_var.set(directory_path)

    ttk.Button(replace_dir_frame, text="参照...", command=browse_replace_directory).pack(side=tk.LEFT)

    # 置換テキスト入力部分
    text_frame = ttk.Frame(replace_frame)
    text_frame.pack(fill=tk.X, pady=(10, 10))

    # 書式保持オプション
    format_option_frame = ttk.Frame(replace_frame)
    format_option_frame.pack(fill=tk.X, pady=(0, 10))

    ttk.Checkbutton(format_option_frame, text="書式を保持する（推奨）", variable=preserve_formatting_var).pack(
        anchor=tk.W)

    ttk.Label(text_frame, text="置換前テキスト:").grid(row=0, column=0, sticky=tk.W, pady=5)
    ttk.Entry(text_frame, textvariable=old_text_var, width=40).grid(row=0, column=1, padx=5, pady=5, sticky=tk.W + tk.E)

    ttk.Label(text_frame, text="置換後テキスト:").grid(row=1, column=0, sticky=tk.W, pady=5)
    ttk.Entry(text_frame, textvariable=new_text_var, width=40).grid(row=1, column=1, padx=5, pady=5, sticky=tk.W + tk.E)

    text_frame.columnconfigure(1, weight=1)

    # 実行ボタン
    replace_button_frame = ttk.Frame(replace_frame)
    replace_button_frame.pack(pady=(10, 20))

    # 進捗バー（置換タブ用）
    replace_progress_frame = ttk.Frame(replace_frame)
    replace_progress_frame.pack(fill=tk.X, pady=(0, 10))
    replace_progress_bar = ttk.Progressbar(replace_progress_frame, orient="horizontal", length=300, mode="determinate")
    replace_progress_bar.pack(fill=tk.X)

    # ステータス表示（置換タブ用）
    replace_status_frame = ttk.Frame(replace_frame)
    replace_status_frame.pack(fill=tk.X)
    replace_status_label = ttk.Label(replace_status_frame, textvariable=status_var, wraplength=500)
    replace_status_label.pack(fill=tk.X)

    def update_replace_progress(value):
        replace_progress_bar["value"] = value
        root.update_idletasks()

    def on_replace():
        directory_path = replace_dir_var.get()
        old_text = old_text_var.get()
        new_text = new_text_var.get()

        if not directory_path:
            messagebox.showerror("エラー", "ディレクトリを選択してください。")
            return

        if not old_text or not new_text:
            messagebox.showerror("エラー", "置換前と置換後のテキストを入力してください。")
            return

        # UIを無効化
        replace_button["state"] = "disabled"
        replace_progress_bar["value"] = 0
        status_var.set("置換処理を開始します...")

        def replace_thread():
            try:
                # 処理を実行
                processed_files, modified_files = replace_text_in_word_files(
                    directory_path,
                    old_text,
                    new_text,
                    preserve_formatting=preserve_formatting_var.get(),
                    progress_callback=update_replace_progress,
                    status_callback=update_status
                )

                # 結果を表示
                messagebox.showinfo("処理完了",
                                    f"処理が完了しました。\n"
                                    f"処理されたファイル数: {processed_files}\n"
                                    f"変更されたファイル数: {modified_files}")

                status_var.set(f"完了しました。処理ファイル数: {processed_files}, 変更ファイル数: {modified_files}")
            except Exception as e:
                messagebox.showerror("エラー", f"処理中にエラーが発生しました: {e}")
                status_var.set(f"エラー: {e}")
            finally:
                # UIを有効化
                root.after(0, lambda: replace_button.config(state="normal"))

        # 別スレッドで処理を実行
        threading.Thread(target=replace_thread, daemon=True).start()

    replace_button = ttk.Button(replace_button_frame, text="置換実行", command=on_replace, width=15)
    replace_button.pack()

    # 終了ボタン
    exit_button_frame = ttk.Frame(root)
    exit_button_frame.pack(pady=10, padx=20)

    def exit_application():
        if messagebox.askokcancel("終了確認", "アプリケーションを終了してもよろしいですか？"):
            root.destroy()

    exit_button = ttk.Button(exit_button_frame, text="終了", command=exit_application, width=15)
    exit_button.pack()

    # メインループを開始
    root.mainloop()


if __name__ == "__main__":
    create_gui()