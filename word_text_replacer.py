# pip install python-docx
import tkinter as tk
from pathlib import Path
from tkinter import filedialog, messagebox

import docx


def replace_text_in_word_files(directory_path, old_text, new_text):
    """
    指定されたディレクトリ内のすべてのWordファイルで特定のテキストを置換します。

    Args:
        directory_path (str): Wordファイルを含むディレクトリのパス
        old_text (str): 置換対象のテキスト
        new_text (str): 置換後のテキスト

    Returns:
        tuple: (処理されたファイル数, 変更されたファイル数)
    """
    # ディレクトリ内のすべてのWordファイルを取得
    word_files = []
    for file in Path(directory_path).glob("**/*.docx"):
        word_files.append(str(file))

    processed_files = len(word_files)
    modified_files = 0

    # 各Wordファイルを処理
    for file_path in word_files:
        try:
            # Wordファイルを開く
            doc = docx.Document(file_path)

            # 置換が行われたかを追跡
            replacement_occurred = False

            # すべての段落内のテキストを置換
            for paragraph in doc.paragraphs:
                if old_text in paragraph.text:
                    replacement_occurred = True
                    text = paragraph.text
                    paragraph.text = text.replace(old_text, new_text)

            # すべての表内のテキストを置換
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if old_text in paragraph.text:
                                replacement_occurred = True
                                text = paragraph.text
                                paragraph.text = text.replace(old_text, new_text)

            # 変更があった場合のみファイルを保存
            if replacement_occurred:
                doc.save(file_path)
                modified_files += 1

        except Exception as e:
            print(f"ファイル '{file_path}' の処理中にエラーが発生しました: {e}")

    return processed_files, modified_files


def select_directory():
    """
    ディレクトリ選択ダイアログを表示し、選択されたディレクトリパスを返します。
    """
    root = tk.Tk()
    root.withdraw()  # メインウィンドウを非表示にする
    directory_path = filedialog.askdirectory(title="Wordファイルを含むディレクトリを選択")
    return directory_path


def create_gui():
    """
    GUIアプリケーションを作成して表示します。
    """

    def on_process():
        directory_path = directory_var.get()
        old_text = old_text_var.get()
        new_text = new_text_var.get()

        if not directory_path:
            messagebox.showerror("エラー", "ディレクトリを選択してください。")
            return

        if not old_text or not new_text:
            messagebox.showerror("エラー", "置換前と置換後のテキストを入力してください。")
            return

        # 処理を実行
        processed_files, modified_files = replace_text_in_word_files(directory_path, old_text, new_text)

        # 結果を表示
        messagebox.showinfo("処理完了",
                            f"処理が完了しました。\n"
                            f"処理されたファイル数: {processed_files}\n"
                            f"変更されたファイル数: {modified_files}")

    def browse_directory():
        directory_path = select_directory()
        if directory_path:
            directory_var.set(directory_path)

    # ルートウィンドウを作成
    root = tk.Tk()
    root.title("Wordファイル置換ツール")
    root.geometry("500x250")
    root.resizable(False, False)

    # 変数を定義
    directory_var = tk.StringVar()
    old_text_var = tk.StringVar()
    new_text_var = tk.StringVar()

    # フレーム作成
    main_frame = tk.Frame(root, padx=20, pady=20)
    main_frame.pack(fill=tk.BOTH, expand=True)

    # ディレクトリ選択部分
    dir_frame = tk.Frame(main_frame)
    dir_frame.pack(fill=tk.X, pady=(0, 10))

    tk.Label(dir_frame, text="ディレクトリ:").pack(side=tk.LEFT)
    tk.Entry(dir_frame, textvariable=directory_var, width=30).pack(side=tk.LEFT, padx=(5, 5), fill=tk.X, expand=True)
    tk.Button(dir_frame, text="参照...", command=browse_directory).pack(side=tk.LEFT)

    # 置換テキスト入力部分
    text_frame = tk.Frame(main_frame)
    text_frame.pack(fill=tk.X, pady=(0, 10))

    tk.Label(text_frame, text="置換前テキスト:").pack(side=tk.LEFT)
    tk.Entry(text_frame, textvariable=old_text_var, width=20).pack(side=tk.LEFT, padx=(5, 10))

    tk.Label(text_frame, text="置換後テキスト:").pack(side=tk.LEFT)
    tk.Entry(text_frame, textvariable=new_text_var, width=20).pack(side=tk.LEFT, padx=(5, 0))

    # 実行ボタン
    button_frame = tk.Frame(main_frame)
    button_frame.pack(pady=(10, 0))

    tk.Button(button_frame, text="実行", command=on_process, width=15, height=2).pack()

    # メインループを開始
    root.mainloop()


if __name__ == "__main__":
    create_gui()
